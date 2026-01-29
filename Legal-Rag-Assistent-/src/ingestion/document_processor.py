from pathlib import Path
from typing import Dict, Any
from loguru import logger

import fitz  # pymupdf
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from config.settings import settings

def extract_text_from_pdf(path: str) -> str:
    doc = fitz.open(path)
    parts = []
    for i in range(len(doc)):
        parts.append(doc.load_page(i).get_text())
    doc.close()
    return "\n".join(parts)


def extract_text_from_docx(path: str) -> str:
    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(parts)


def extract_text_from_txt(path: str) -> str:
    p = Path(path)
    try:
        return p.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return p.read_text(encoding="latin-1")


def process_file(path: str) -> Dict[str, Any]:
    suffix = Path(path).suffix.lower()

    if suffix == ".pdf":
        text = extract_text_from_pdf(path)
    elif suffix in [".docx", ".doc"]:
        text = extract_text_from_docx(path)
    elif suffix == ".txt":
        text = extract_text_from_txt(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    # minimal cleaning
    text = " ".join(text.split())

    meta = {
        "source": Path(path).name,
        "path": str(Path(path).resolve()),
        "type": suffix.replace(".", "").upper(),
    }

    logger.info(f"Processed {path} (chars={len(text)})")
    return {"text": text, "metadata": meta}
def load_documents(target_dir: str = None) -> list[Document]:
    """
    Loads all PDFs/DOCX/TXT from settings.DOCS_DIR and returns LangChain Documents.
    """
    path_str = target_dir if target_dir else settings.DOCS_DIR
    docs_dir = Path(path_str)
    if not docs_dir.exists():
        logger.warning(f"DOCS_DIR not found: {docs_dir}")
        return []

    files = [p for p in docs_dir.iterdir() if p.is_file()]
    if not files:
        logger.warning(f"No files in: {docs_dir}")
        return []

    docs: list[Document] = []

    for f in files:
        suffix = f.suffix.lower()

        # PDFs: page-wise loader (best for citations)
        if suffix == ".pdf":
            loader = PyPDFLoader(str(f))
            pages = loader.load()
            for p in pages:
                p.metadata = p.metadata or {}
                p.metadata["source"] = f.name
            docs.extend(pages)
            continue

        # DOCX/TXT: treat as single document
        out = process_file(str(f))
        docs.append(Document(page_content=out["text"], metadata=out["metadata"]))

    logger.info(f"Loaded {len(docs)} documents/pages from {docs_dir}")
    return docs


def split_documents(docs: list[Document]) -> list[Document]:
    """
    Splits documents into chunks using RecursiveCharacterTextSplitter.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
    )
    chunks = splitter.split_documents(docs)

    for i, c in enumerate(chunks):
        c.metadata = c.metadata or {}
        c.metadata["chunk"] = c.metadata.get("chunk", i)

    logger.info(f"Split into {len(chunks)} chunks")
    return chunks


if __name__ == "__main__":
    # run this file alone for testing
    sample_dir = Path("data/uploads")
    files = [p for p in sample_dir.iterdir() if p.is_file()]
    if not files:
        print("No files in data/uploads")
    else:
        out = process_file(str(files[0]))
        print("OK:", out["metadata"], "text_chars=", len(out["text"]))
